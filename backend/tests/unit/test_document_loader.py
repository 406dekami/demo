"""
文档加载器测试
"""
import pytest
import os
import tempfile
from backend.app.rag.document_loader import DocumentLoader


class TestDocumentLoader:
    """测试 DocumentLoader 类"""
    
    def test_load_txt_file(self):
        """测试加载 TXT 文件"""
        # 创建临时 TXT 文件
        content = "这是第一行。\n这是第二行。\n这是第三行。"
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write(content)
            temp_path = f.name
        
        try:
            loader = DocumentLoader()
            result = loader.load_document(temp_path)
            
            assert len(result) > 0
            assert "text" in result[0]
            assert "source" in result[0]
            assert "这是第一行" in result[0]["text"]
        finally:
            # 清理临时文件
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    def test_load_md_file(self):
        """测试加载 Markdown 文件"""
        content = "# 标题\n\n这是内容。\n\n## 子标题\n\n更多。"
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False, encoding='utf-8') as f:
            f.write(content)
            temp_path = f.name
        
        try:
            loader = DocumentLoader()
            result = loader.load_document(temp_path)
            
            assert len(result) > 0
            assert "# 标题" in result[0]["text"]
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    def test_load_docx_file(self):
        """测试加载 Word 文件"""
        from docx import Document as DocxDocument
        
        # 创建临时 Word 文件
        temp_path = tempfile.mktemp(suffix='.docx')
        doc = DocxDocument()
        doc.add_paragraph("第一段内容。")
        doc.add_paragraph("第二段内容。")
        doc.save(temp_path)
        
        try:
            loader = DocumentLoader()
            result = loader.load_document(temp_path)
            
            assert len(result) > 0
            assert "第一段内容" in result[0]["text"]
            assert "第二段内容" in result[0]["text"]
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    def test_load_pdf_file(self):
        """测试加载 PDF 文件"""
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import tempfile
        import os
        
        # 创建临时 PDF 文件
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            temp_path = f.name
        
        try:
            # 使用 reportlab 创建 PDF，注册并使用中文字体
            c = canvas.Canvas(temp_path)
            
            # 尝试注册系统字体（Windows 常见中文字体）
            font_found = False
            font_paths = [
                r"C:\Windows\Fonts\simhei.ttf",  # 黑体
                r"C:\Windows\Fonts\simsun.ttc",  # 宋体
                r"C:\Windows\Fonts\msyh.ttf",    # 微软雅黑
            ]
            
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                        c.setFont('ChineseFont', 12)
                        font_found = True
                        break
                    except Exception:
                        continue
            
            if not font_found:
                # 如果没有中文字体，使用默认字体（英文测试）
                c.setFont('Helvetica', 12)
                c.drawString(100, 750, "PDF Test Content Line 1")
                c.drawString(100, 730, "PDF Test Content Line 2")
                print("使用默认字体")
            else:
                # 使用中文字体
                c.drawString(100, 750, "这是 PDF 测试内容第一行")
                c.drawString(100, 730, "这是 PDF 测试内容第二行")
                print("使用中文字体")
            
            c.save()
            
            # 测试加载
            loader = DocumentLoader()
            result = loader.load_document(temp_path)
            
            assert len(result) > 0
            assert "text" in result[0]
            assert "page" in result[0]
            assert "source" in result[0]
            # 验证内容是否被正确提取
            full_text = " ".join([chunk["text"] for chunk in result])
            assert len(full_text.strip()) > 0
            
            # 根据使用的字体验证内容
            if font_found:
                assert "测试内容" in full_text or "PDF" in full_text
            else:
                assert "Test Content" in full_text
            print(f"full_txt:{full_text}")
        except ImportError:
            pytest.skip("reportlab 未安装，无法创建 PDF")
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    def test_load_unsupported_format(self):
        """测试加载不支持的文件格式"""
        loader = DocumentLoader()
        
        # 创建临时文件
        with tempfile.NamedTemporaryFile(suffix='.xyz', delete=False) as f:
            f.write(b"test")
            temp_path = f.name
        
        try:
            with pytest.raises(ValueError, match="不支持的文件格式"):
                loader.load_document(temp_path)
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    def test_load_empty_file(self):
        """测试加载空文件"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("")
            temp_path = f.name
        
        try:
            loader = DocumentLoader()
            result = loader.load_document(temp_path)
            
            # 空文件可能返回空列表或包含空文本的列表
            assert isinstance(result, list)
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
    
    def test_load_file_with_metadata(self):
        """测试加载文件包含元数据"""
        content = "测试内容"
        
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write(content)
            temp_path = f.name
        
        try:
            loader = DocumentLoader()
            result = loader.load_document(temp_path)
            
            assert len(result) > 0
            assert "source" in result[0]
            assert os.path.basename(temp_path) in result[0]["source"]
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
