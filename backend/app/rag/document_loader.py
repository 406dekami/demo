"""
文档加载与解析模块
支持多种格式的文档加载和解析
"""
import logging
import os
from typing import List, Dict, Any

import pdfplumber
from docx import Document as DocxDocument


class DocumentLoader:
    """
    文档加载器，支持多种格式的文档解析
    """
    
    def __init__(self):
        self.supported_types = ['pdf', 'txt', 'md', 'docx', 'doc']
        
    def load_document(self, file_path: str) -> List[Dict[str, Any]]:
        """
        加载并解析文档，返回文本片段列表
        
        Args:
            file_path: 文档文件路径
            
        Returns:
            包含文本内容和元信息的字典列表
        """
        file_ext = os.path.splitext(file_path)[1][1:].lower()
        
        if file_ext == 'pdf':
            return self._load_pdf(file_path)
        elif file_ext == 'txt':
            return self._load_txt(file_path)
        elif file_ext == 'md':
            return self._load_md(file_path)
        elif file_ext == 'docx':
            return self._load_docx(file_path)
        elif file_ext == 'doc':
            return self._load_doc(file_path)
        else:
            raise ValueError(f"不支持的文件格式：{file_ext}")

    def _load_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        """
        解析 PDF 文档
        """
        content = []
        try:
            logging.info(f"开始解析 PDF: {file_path}")
            with pdfplumber.open(file_path) as pdf:
                logging.info(f"PDF 页数：{len(pdf.pages)}")
                for page_num, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    logging.info(f"第 {page_num + 1} 页提取文本长度：{len(text) if text else 0}")
                    if text and text.strip():
                        content.append({
                            "text": text,
                            "page": page_num + 1,
                            "source": f"{os.path.basename(file_path)}_page_{page_num + 1}"
                        })
        except Exception as e:
            logging.error(f"解析 PDF 文件失败 {file_path}: {e}")
            raise
            
        logging.info(f"PDF 解析完成，共提取 {len(content)} 个文本片段")
        return content
    
    def _load_txt(self, file_path: str) -> List[Dict[str, Any]]:
        """
        解析纯文本文件
        """
        content = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                if text and text.strip():
                    content.append({
                        "text": text,
                        "source": os.path.basename(file_path)
                    })
        except Exception as e:
            logging.error(f"解析TXT文件失败 {file_path}: {e}")
            raise
        
        return content
    
    def _load_md(self, file_path: str) -> List[Dict[str, Any]]:
        """
        解析Markdown文件
        """
        content = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                if text and text.strip():
                    content.append({
                        "text": text,
                        "source": os.path.basename(file_path)
                    })
        except Exception as e:
            logging.error(f"解析Markdown文件失败 {file_path}: {e}")
            raise
        
        return content
    
    def _load_docx(self, file_path: str) -> List[Dict[str, Any]]:
        """
        解析 Word 文档
        """
        content = []
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            if text and text.strip():
                content.append({
                    "text": text,
                    "source": os.path.basename(file_path)
                })
        except Exception as e:
            logging.error(f"解析 Word 文件失败 {file_path}: {e}")
            raise
            
        return content
        
    def _load_doc(self, file_path: str) -> List[Dict[str, Any]]:
        """
        解析旧版 Word 文档（.doc）
        注意：.doc 是二进制格式，需要使用专用库或转换工具
        """
        content = []
        try:
            # 方案 1：使用 pywin32 COM 接口（Windows 环境，最可靠）
            try:
                import win32com.client
                import tempfile
                
                logging.info("使用 COM 接口解析 .doc 文件...")
                word = None
                doc = None
                
                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    word.DisplayAlerts = 0  # 0 = wdAlertsNone
                    
                    # 打开 .doc 文件
                    doc = word.Documents.Open(os.path.abspath(file_path))
                    
                    # 另存为 .docx 临时文件
                    temp_docx = tempfile.mktemp(suffix=".docx")
                    doc.SaveAs2(temp_docx, FileFormat=16)  # 16 = docx format
                    
                    if doc:
                        doc.Close(SaveChanges=False)
                    if word:
                        word.Quit()
                    
                    # 强制释放 COM 对象
                    import gc
                    del doc
                    del word
                    gc.collect()
                    
                    # 用 python-docx 解析临时文件
                    temp_doc = DocxDocument(temp_docx)
                    text = ""
                    for paragraph in temp_doc.paragraphs:
                        text += paragraph.text + "\n"
                    
                    # 清理临时文件
                    if os.path.exists(temp_docx):
                        os.remove(temp_docx)
                    
                    if text and text.strip():
                        content.append({
                            "text": text,
                            "source": os.path.basename(file_path)
                        })
                        logging.info(f"COM 接口解析成功，提取文本长度：{len(text)}")
                        return content
                except Exception as e:
                    # 确保在出错时也清理 COM 资源
                    try:
                        if doc:
                            doc.Close(SaveChanges=False)
                    except:
                        pass
                    try:
                        if word:
                            word.Quit()
                    except:
                        pass
                    import gc
                    gc.collect()
                    raise e
                    
            except ImportError:
                logging.info("pywin32 未安装，跳过 COM 接口解析")
            except Exception as e:
                logging.warning(f"COM 接口解析失败：{e}")
            
            # 方案 2：尝试 macOS textutil（macOS 自带工具）
            import subprocess
            
            try:
                import tempfile
                temp_docx = tempfile.mktemp(suffix=".docx")
                result = subprocess.run(
                    ['textutil', '-convert', 'docx', '-output', temp_docx, file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and os.path.exists(temp_docx):
                    # 用 python-docx 解析转换后的文件
                    temp_doc = DocxDocument(temp_docx)
                    text = ""
                    for paragraph in temp_doc.paragraphs:
                        text += paragraph.text + "\n"
                    
                    # 清理临时文件
                    os.remove(temp_docx)
                    
                    if text.strip():
                        content.append({
                            "text": text,
                            "source": os.path.basename(file_path)
                        })
                        logging.info(f"textutil 转换解析成功，提取文本长度：{len(text)}")
                        return content
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass
            
            # 方案 3：尝试 antiword 或 catdoc（Linux/Mac）
            try:
                result = subprocess.run(
                    ['antiword', file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    content.append({
                        "text": result.stdout,
                        "source": os.path.basename(file_path)
                    })
                    return content
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass
            
            try:
                result = subprocess.run(
                    ['catdoc', file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout.strip():
                    content.append({
                        "text": result.stdout,
                        "source": os.path.basename(file_path)
                    })
                    return content
            except (FileNotFoundError, subprocess.TimeoutExpired):
                pass
            
            # 方案 3：使用 olefile（纯 Python，效果有限）
            try:
                import olefile
                ole = olefile.OleFileIO(file_path)
                text_streams = []
                
                for stream in ole.listdir():
                    if 'WordDocument' in stream or '1Table' in stream:
                        continue
                    try:
                        data = ole.openstream(stream).read()
                        try:
                            text = data.decode('utf-8', errors='ignore')
                            if text.strip():
                                text_streams.append(text)
                        except:
                            pass
                    except:
                        pass
                
                if text_streams:
                    content.append({
                        "text": "\n".join(text_streams),
                        "source": os.path.basename(file_path)
                    })
            except ImportError:
                logging.warning("olefile 未安装，无法解析 .doc 文件")
            except Exception as e:
                logging.warning(f"olefile 解析失败：{e}")
            
            # 如果所有方案都失败，给出明确提示
            if not content:
                error_msg = (
                    f"无法解析 .doc 文件：{os.path.basename(file_path)}\n"
                    f"建议：\n"
                    f"1. 将 .doc 文件转换为 .docx 或 .pdf 后重新上传（推荐）\n"
                    f"2. 在 Windows 上安装 pywin32：pip install pywin32\n"
                    f"3. Linux 环境安装 antiword：sudo apt-get install antiword"
                )
                logging.error(error_msg)
                raise ValueError(error_msg)
                        
        except Exception as e:
            logging.error(f"解析 DOC 文件失败 {file_path}: {e}")
            raise
        
        return content